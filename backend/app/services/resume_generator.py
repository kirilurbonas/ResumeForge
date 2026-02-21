"""Resume generator service for creating DOC and PDF files."""

import io
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

from app.models.resume_model import Resume
from app.services.template_engine import TemplateEngine


class ResumeGenerator:
    """Generate resumes in DOC and PDF formats."""
    
    def __init__(self):
        self.template_engine = TemplateEngine()
    
    def generate_doc(self, resume: Resume, template_id: str = "modern", custom_template: Optional[dict] = None) -> bytes:
        """
        Generate resume as DOCX file.
        
        Args:
            resume: Resume object
            template_id: Template ID to use
            custom_template: Optional custom template dict to override defaults
            
        Returns:
            DOCX file as bytes
        """
        template = self.template_engine.get_template(template_id)
        if not template:
            template = self.template_engine.get_template("modern")  # Fallback
        
        # Merge custom template if provided
        if custom_template:
            template = {**template, **custom_template}
        
        # Create document
        doc = Document()
        
        # Apply template settings
        self._apply_doc_template(doc, template, resume)
        
        # Convert to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.read()
    
    def generate_pdf(self, resume: Resume, template_id: str = "modern", custom_template: Optional[dict] = None) -> bytes:
        """
        Generate resume as PDF file.
        
        Args:
            resume: Resume object
            template_id: Template ID to use
            
        Returns:
            PDF file as bytes
        """
        template = self.template_engine.get_template(template_id)
        if not template:
            template = self.template_engine.get_template("modern")  # Fallback
        
        # Merge custom template if provided
        if custom_template:
            template = {**template, **custom_template}
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        # Build content
        story = []
        self._build_pdf_content(story, template, resume)
        
        # Build PDF
        doc.build(story)
        
        buffer.seek(0)
        return buffer.read()
    
    def _apply_doc_template(self, doc: Document, template: dict, resume: Resume):
        """Apply template styling to DOCX document."""
        fonts = template.get('fonts', {})
        colors = template.get('colors', {})
        spacing = template.get('spacing', {})
        style_config = template.get('style', {})
        
        # Header section
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if style_config.get('header_alignment') == 'center' else WD_ALIGN_PARAGRAPH.LEFT
        
        # Name
        name_run = header_para.add_run(resume.contact_info.name)
        name_run.font.size = Pt(fonts.get('heading_size', 18))
        name_run.font.color.rgb = RGBColor.from_string(colors.get('primary', '#000000').lstrip('#'))
        name_run.bold = True
        
        # Contact info
        contact_info = []
        if resume.contact_info.email:
            contact_info.append(resume.contact_info.email)
        if resume.contact_info.phone:
            contact_info.append(resume.contact_info.phone)
        if resume.contact_info.location:
            contact_info.append(resume.contact_info.location)
        if resume.contact_info.linkedin:
            contact_info.append(resume.contact_info.linkedin)
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = header_para.alignment
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(fonts.get('body_size', 11))
        
        # Summary
        if resume.summary:
            doc.add_paragraph()  # Spacing
            summary_heading = doc.add_paragraph('SUMMARY' if style_config.get('section_headers_uppercase') else 'Summary')
            summary_heading.runs[0].bold = True
            summary_heading.runs[0].font.size = Pt(fonts.get('heading_size', 16))
            
            summary_para = doc.add_paragraph(resume.summary)
            summary_para.runs[0].font.size = Pt(fonts.get('body_size', 11))
        
        # Experience
        if resume.experience:
            doc.add_paragraph()  # Spacing
            exp_heading = doc.add_paragraph('EXPERIENCE' if style_config.get('section_headers_uppercase') else 'Experience')
            exp_heading.runs[0].bold = True
            exp_heading.runs[0].font.size = Pt(fonts.get('heading_size', 16))
            
            for exp in resume.experience:
                # Position and company
                exp_header = doc.add_paragraph()
                exp_header.add_run(f"{exp.position}").bold = True
                exp_header.add_run(f" - {exp.company}")
                exp_header.runs[0].font.size = Pt(fonts.get('body_size', 11))
                
                # Dates
                date_str = f"{exp.start_date} - {exp.end_date if exp.end_date else 'Present'}"
                date_para = doc.add_paragraph(date_str)
                date_para.runs[0].font.size = Pt(fonts.get('body_size', 10))
                date_para.runs[0].italic = True
                
                # Description
                for desc in exp.description:
                    desc_para = doc.add_paragraph(desc, style='List Bullet' if style_config.get('use_bullets') else None)
                    desc_para.runs[0].font.size = Pt(fonts.get('body_size', 11))
        
        # Education
        if resume.education:
            doc.add_paragraph()  # Spacing
            edu_heading = doc.add_paragraph('EDUCATION' if style_config.get('section_headers_uppercase') else 'Education')
            edu_heading.runs[0].bold = True
            edu_heading.runs[0].font.size = Pt(fonts.get('heading_size', 16))
            
            for edu in resume.education:
                edu_para = doc.add_paragraph()
                edu_para.add_run(f"{edu.degree}").bold = True
                if edu.field_of_study:
                    edu_para.add_run(f" in {edu.field_of_study}")
                edu_para.add_run(f", {edu.institution}")
                edu_para.runs[0].font.size = Pt(fonts.get('body_size', 11))
                
                date_str = f"{edu.start_date} - {edu.end_date if edu.end_date else 'Present'}"
                date_para = doc.add_paragraph(date_str)
                date_para.runs[0].font.size = Pt(fonts.get('body_size', 10))
                date_para.runs[0].italic = True
        
        # Skills
        if resume.skills:
            doc.add_paragraph()  # Spacing
            skills_heading = doc.add_paragraph('SKILLS' if style_config.get('section_headers_uppercase') else 'Skills')
            skills_heading.runs[0].bold = True
            skills_heading.runs[0].font.size = Pt(fonts.get('heading_size', 16))
            
            skill_names = [skill.name for skill in resume.skills]
            skills_para = doc.add_paragraph(', '.join(skill_names))
            skills_para.runs[0].font.size = Pt(fonts.get('body_size', 11))
        
        # Certifications
        if resume.certifications:
            doc.add_paragraph()  # Spacing
            cert_heading = doc.add_paragraph('CERTIFICATIONS' if style_config.get('section_headers_uppercase') else 'Certifications')
            cert_heading.runs[0].bold = True
            cert_heading.runs[0].font.size = Pt(fonts.get('heading_size', 16))
            
            for cert in resume.certifications:
                cert_para = doc.add_paragraph(f"{cert.name} - {cert.issuer} ({cert.date})")
                cert_para.runs[0].font.size = Pt(fonts.get('body_size', 11))
    
    def _build_pdf_content(self, story: list, template: dict, resume: Resume):
        """Build PDF content using reportlab."""
        styles = getSampleStyleSheet()
        
        fonts = template.get('fonts', {})
        colors = template.get('colors', {})
        style_config = template.get('style', {})
        
        # Custom styles
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=fonts.get('heading_size', 16),
            textColor=colors.get('primary', '#000000'),
            spaceAfter=6,
            alignment=TA_CENTER if style_config.get('header_alignment') == 'center' else TA_LEFT
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=fonts.get('body_size', 11),
            spaceAfter=4
        )
        
        # Header
        name_para = Paragraph(resume.contact_info.name, heading_style)
        story.append(name_para)
        story.append(Spacer(1, 0.1*inch))
        
        # Contact info
        contact_info = []
        if resume.contact_info.email:
            contact_info.append(resume.contact_info.email)
        if resume.contact_info.phone:
            contact_info.append(resume.contact_info.phone)
        if resume.contact_info.location:
            contact_info.append(resume.contact_info.location)
        
        if contact_info:
            contact_para = Paragraph(' | '.join(contact_info), body_style)
            story.append(contact_para)
            story.append(Spacer(1, 0.2*inch))
        
        # Summary
        if resume.summary:
            summary_heading = Paragraph('SUMMARY' if style_config.get('section_headers_uppercase') else 'Summary', styles['Heading2'])
            story.append(summary_heading)
            summary_para = Paragraph(resume.summary, body_style)
            story.append(summary_para)
            story.append(Spacer(1, 0.15*inch))
        
        # Experience
        if resume.experience:
            exp_heading = Paragraph('EXPERIENCE' if style_config.get('section_headers_uppercase') else 'Experience', styles['Heading2'])
            story.append(exp_heading)
            
            for exp in resume.experience:
                exp_text = f"<b>{exp.position}</b> - {exp.company}"
                exp_para = Paragraph(exp_text, body_style)
                story.append(exp_para)
                
                date_text = f"{exp.start_date} - {exp.end_date if exp.end_date else 'Present'}"
                date_para = Paragraph(date_text, body_style)
                story.append(date_para)
                
                for desc in exp.description:
                    desc_text = f"• {desc}" if style_config.get('use_bullets') else desc
                    desc_para = Paragraph(desc_text, body_style)
                    story.append(desc_para)
                
                story.append(Spacer(1, 0.1*inch))
        
        # Education
        if resume.education:
            edu_heading = Paragraph('EDUCATION' if style_config.get('section_headers_uppercase') else 'Education', styles['Heading2'])
            story.append(edu_heading)
            
            for edu in resume.education:
                edu_text = f"<b>{edu.degree}</b>"
                if edu.field_of_study:
                    edu_text += f" in {edu.field_of_study}"
                edu_text += f", {edu.institution}"
                edu_para = Paragraph(edu_text, body_style)
                story.append(edu_para)
                
                date_text = f"{edu.start_date} - {edu.end_date if edu.end_date else 'Present'}"
                date_para = Paragraph(date_text, body_style)
                story.append(date_para)
                story.append(Spacer(1, 0.1*inch))
        
        # Skills
        if resume.skills:
            skills_heading = Paragraph('SKILLS' if style_config.get('section_headers_uppercase') else 'Skills', styles['Heading2'])
            story.append(skills_heading)
            
            skill_names = [skill.name for skill in resume.skills]
            skills_text = ', '.join(skill_names)
            skills_para = Paragraph(skills_text, body_style)
            story.append(skills_para)
            story.append(Spacer(1, 0.15*inch))
        
        # Certifications
        if resume.certifications:
            cert_heading = Paragraph('CERTIFICATIONS' if style_config.get('section_headers_uppercase') else 'Certifications', styles['Heading2'])
            story.append(cert_heading)
            
            for cert in resume.certifications:
                cert_text = f"{cert.name} - {cert.issuer} ({cert.date})"
                cert_para = Paragraph(cert_text, body_style)
                story.append(cert_para)
