import io

import pytest
from docx import Document

from app.services.resume_parser import ResumeParser


def test_parser_rejects_unknown_extension():
    parser = ResumeParser()
    with pytest.raises(ValueError, match="Unsupported"):
        parser.parse(b"hello", "file.xyz")


def test_parser_docx_minimal():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Alex Smith")
    doc.add_paragraph("alex@example.com")
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, AWS")
    doc.save(buf)
    buf.seek(0)

    parser = ResumeParser()
    resume = parser.parse(buf.getvalue(), "cv.docx")

    assert resume.contact_info.email == "alex@example.com"
    assert resume.filename == "cv.docx"
    assert resume.id
