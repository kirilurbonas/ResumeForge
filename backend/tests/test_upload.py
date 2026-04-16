import io

import pytest
from docx import Document


def _minimal_docx_bytes() -> bytes:
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane.doe@example.com")
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Engineer at Acme Corp")
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def test_upload_rejects_bad_extension(client, auth_headers):
    r = client.post(
        "/api/resume/upload",
        headers=auth_headers,
        files={"file": ("resume.txt", b"hello", "text/plain")},
    )
    assert r.status_code == 400
    assert "Unsupported" in r.json().get("detail", "")


def test_upload_rejects_empty_file(client, auth_headers):
    r = client.post(
        "/api/resume/upload",
        headers=auth_headers,
        files={"file": ("empty.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 400


def test_upload_docx_success(client, auth_headers):
    content = _minimal_docx_bytes()
    r = client.post(
        "/api/resume/upload",
        headers=auth_headers,
        files={
            "file": (
                "resume.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert "id" in data
    assert data["filename"] == "resume.docx"


def test_upload_rejects_oversized(monkeypatch, client, auth_headers):
    from app.api import routes as routes_module

    monkeypatch.setattr(routes_module, "MAX_FILE_SIZE", 50)
    big = b"x" * 100
    r = client.post(
        "/api/resume/upload",
        headers=auth_headers,
        files={
            "file": (
                "big.docx",
                big,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 400
    assert "exceeds" in r.json().get("detail", "").lower()
