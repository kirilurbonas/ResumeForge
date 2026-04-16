"""Duplicate resume and list search (q) API."""

from docx import Document
import io


def _docx_bytes():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Searchable NameX")
    doc.add_paragraph("searchname@example.com")
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def test_duplicate_resume(client, auth_headers):
    up = client.post(
        "/api/resume/upload",
        headers=auth_headers,
        files={
            "file": (
                "original.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert up.status_code == 200
    rid = up.json()["id"]

    dup = client.post(
        f"/api/resume/{rid}/duplicate",
        headers=auth_headers,
    )
    assert dup.status_code == 200, dup.text
    assert dup.json()["id"] != rid
    assert "copy" in dup.json()["filename"].lower()

    lst = client.get("/api/resumes", headers=auth_headers)
    assert lst.status_code == 200
    assert lst.json()["count"] >= 2


def test_list_resumes_search_q(client, auth_headers):
    up = client.post(
        "/api/resume/upload",
        headers=auth_headers,
        files={
            "file": (
                "unique_alpha_beta.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert up.status_code == 200

    r = client.get(
        "/api/resumes",
        headers=auth_headers,
        params={"q": "unique_alpha"},
    )
    assert r.status_code == 200
    names = [x["filename"] for x in r.json().get("resumes", [])]
    assert any("unique_alpha" in n for n in names)
